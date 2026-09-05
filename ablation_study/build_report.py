#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
R = ROOT / "results"
FIG = R / "figures"
OUT = ROOT / "reports" / "ablation_results_qwen_api.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "163A5F"
INK = "0B2545"
MUTED = "5D6875"
LIGHT = "F2F4F7"
PALE = "E8EEF5"
GOLD = "FFF4D6"
WHITE = "FFFFFF"
GRID = "B8C2CC"
FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)
    pPr.append(OxmlElement("w:contextualSpacing"))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(TABLE_WIDTH))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(TABLE_INDENT))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(w / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_table(table, widths, numeric_cols=()):
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    for ri, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
        if ri == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            trPr.append(repeat)
        for ci, cell in enumerate(row.cells):
            if ri == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT)
                cell._tc.get_or_add_tcPr().append(shd)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, 8.5, bold=(ri == 0), color=INK if ri == 0 else None)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, 9, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True, color=INK)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name = FONT_LATIN
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    cap = doc.styles["Caption"]
    cap.font.name = FONT_LATIN
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor.from_string(MUTED)

    header = sec.header.paragraphs[0]
    header.text = "QWEN API ABLATION STUDY  /  RECONSTRUCTED EMPIRICAL REPORT"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_font(run, 8, bold=True, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.add_run("Independent experiment report   •   ")
    for run in footer.runs:
        set_font(run, 8, color=MUTED)
    add_page_number(footer)


def pct(x):
    return f"{100*x:.1f}%"


def pformat(x):
    return "<0.001" if x < .001 else f"{x:.3f}"


def main():
    cfg = json.loads((R / "protocol_config.json").read_text(encoding="utf-8"))
    s = pd.read_csv(R / "condition_summary.csv").set_index("condition")
    c = pd.read_csv(R / "paired_comparisons.csv")

    doc = Document()
    setup_styles(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Qwen API 消融实验结果报告")
    set_font(r, 24, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("面向情感敏感型需求获取多智能体框架的重建性端到端评估")
    set_font(r, 12, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    vals = [
        ("实验模型", "qwen3.7-plus-2026-05-26（固定快照，API）"),
        ("测试规模", "30 个冻结案例 × 6 个条件；210 次调用全部成功"),
        ("主分析", "150 次唯一流水线调用；60 次重复生成仅保留为审计记录"),
        ("报告日期", "2026-08-10（Asia/Shanghai）"),
    ]
    for i, (a, b) in enumerate(vals):
        meta.cell(i, 0).text = a
        meta.cell(i, 1).text = b
    style_table(meta, [2100, 7260])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, PALE)
    r = p.add_run("核心结论　")
    set_font(r, 11, bold=True, color=DARK)
    r = p.add_run(
        "Validator 是唯一呈现明确工程增益方向的组件：在共享同一份 Structurer 原始输出时，"
        "Traceability 和 Full compliance 均由 83.3% 提升至 100.0%（+16.7 个百分点）。"
        "但在 30 个案例和 Holm 校正后未达到 α=0.05。Affect 与 Need/RAG 未显示显著的语义增益，"
        "因此本次实验支持“工程控制与可审计性改善”，不支持“语义理解显著提升”。"
    )
    set_font(r, 11, color=INK)

    add_heading(doc, "1. 实验性质与复现边界", 1)
    add_body(doc, "本报告记录一次独立、可审计的重建性消融实验。工作区仅提供论文 papernew.docx，未提供原始代码、原30例ID、提示词、预测文件或配套 workbook。因此，本次结果不能被称为原论文表4.7的逐字复现，而是按照论文声明的公开数据、模块职责、随机种子和比较逻辑重新实现的实验。")
    add_body(doc, "数据集与论文描述精确匹配：原始文件含1,482行、781个唯一 questionID；按出现次数不少于20筛得18个主题和738个可用问题。模型由论文中的本地0.5B模型改为用户指定的 Qwen API。该变化显著提高了结构化输出能力，也导致 JSON 和 SRS 指标出现天花板效应。")
    p = add_body(doc, "解释原则：本报告将结构化合规、证据绑定和格式可审计性视为工程指标；Topic-alignment F1作为弱语义代理；不把任何结果解释为临床正确性、心理诊断能力或真实世界部署有效性。", bold_lead="解释原则：")
    shade_paragraph(p, GOLD)

    add_heading(doc, "2. 实验问题与条件", 1)
    add_body(doc, "实验回答四个问题：完整框架是否优于单次LLM；Validator是否改善证据绑定；Affect Interpreter是否增加主题或证据相关性；Need-topic与RAG是否带来可测量的增益。")
    cond = doc.add_table(rows=1, cols=6)
    for j, h in enumerate(["条件", "说明", "Affect", "Need", "RAG", "Validator"]):
        cond.cell(0, j).text = h
    condition_rows = [
        ("B0", "Single-pass LLM", "—", "—", "—", "—"),
        ("B1", "Single LLM + RAG", "—", "✓", "✓", "—"),
        ("A1", "Without Affect", "—", "✓", "✓", "✓"),
        ("A2", "Without Validator", "✓", "✓", "✓", "—"),
        ("A4", "Without Need/RAG", "✓", "—", "—", "✓"),
        ("A3", "Full Agent", "✓", "✓", "✓", "✓"),
    ]
    for row in condition_rows:
        cells = cond.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = value
    style_table(cond, [720, 3240, 1080, 1080, 1080, 2160], numeric_cols=(2,3,4,5))
    add_caption(doc, "表1　消融实验条件。A4为本次补充条件，用于单独检验 Need/RAG。")

    add_heading(doc, "3. 实验设置", 1)
    add_heading(doc, "3.1 数据划分与辅助模块", 2)
    add_body(doc, "使用 seed=42 对738个合格问题进行确定性打乱，得到479个训练样本、111个开发样本和148个测试样本；从测试部分构造90例端到端池，冻结最后30例作为本次比较集。训练集与最终案例均按 questionID 去重。")
    add_body(doc, "Need-topic 模块使用字符级 TF-IDF（1–4 gram，最多60,000特征）和 One-vs-Rest LinearSVC，从18个主题中输出Top-3。RAG以同一向量空间进行余弦检索，每例返回Top-2历史主题模式。检索文本仅作为路由提示，不允许作为需求证据。")
    add_heading(doc, "3.2 Qwen API 固定设置", 2)
    api = doc.add_table(rows=1, cols=2)
    api.cell(0,0).text, api.cell(0,1).text = "项目", "冻结值"
    api_rows = [
        ("模型", "qwen3.7-plus-2026-05-26"),
        ("地域/接口", "China (Beijing), OpenAI-compatible Chat Completions"),
        ("生成模式", "enable_thinking=false; temperature=0"),
        ("结构化输出", "response_format=json_object"),
        ("每例输出", "1–2条SRS需求；必须附 source_turn_ids 与 evidence_quotes"),
        ("API成功率", "210/210；无失败调用"),
    ]
    for a,b in api_rows:
        row=api.add_row().cells; row[0].text=a; row[1].text=b
    style_table(api, [2400, 6960])

    add_heading(doc, "3.3 严格控制变量处理", 2)
    add_body(doc, "API在 temperature=0 时仍可能因服务端实现产生微小非确定性。为使 Validator 消融只改变验证步骤，A3与A2共享同一份原始生成，分别经过/不经过确定性 Validator；A1与B1同样共享无 Affect 的原始生成。额外产生的A1、A3独立生成各30份被保留在 api_responses.jsonl 中，但排除于主分析。")
    add_body(doc, "Validator仅执行可审计的工程操作：将 evidence quote 规范为所引用 evidence ID 对应的原文、规范闭集字段、在语义陈述缺失时拒绝输出。它不新增需求语义，也不把检索示例作为证据。")

    add_heading(doc, "4. 指标与统计方法", 1)
    metrics = doc.add_table(rows=1, cols=3)
    for j,h in enumerate(["指标", "操作定义", "性质"]): metrics.cell(0,j).text=h
    metric_rows = [
        ("JSON parse", "最终记录可解析为JSON对象", "结构"),
        ("SRS form", "所有接受需求以“The system shall”开头", "结构"),
        ("Traceability", "source ID有效且evidence quote为精确原文", "工程"),
        ("Full compliance", "JSON、SRS、证据绑定和闭集类型同时满足", "工程"),
        ("Topic-alignment F1", "输出need labels与数据集主题的逐例F1", "弱语义代理"),
        ("Lexical grounding proxy", "需求内容词与对话内容词的覆盖比例", "自动代理，非人工有效性"),
        ("Latency / tokens", "每个端到端条件的API时间与token", "成本"),
    ]
    for rowv in metric_rows:
        row=metrics.add_row().cells
        for j,v in enumerate(rowv): row[j].text=v
    style_table(metrics, [2100, 5460, 1800])
    add_body(doc, "所有结果以案例为配对单位。二元指标使用 exact McNemar 检验；连续指标使用 Wilcoxon signed-rank；差值95%区间使用2,000次案例级bootstrap。Holm校正在每个指标的5项“完整系统 vs 其他条件”比较内执行。显著性阈值为0.05。")

    add_heading(doc, "5. 主要结果", 1)
    summary = doc.add_table(rows=1, cols=7)
    for j,h in enumerate(["条件", "JSON", "SRS", "Trace", "Full", "Topic F1", "修复/例"]): summary.cell(0,j).text=h
    order = ["B0_single_llm", "B1_single_llm_rag", "A1_without_affect", "A2_without_validator", "A4_without_need_rag", "A3_full_agent"]
    codes = dict(zip(order,["B0","B1","A1","A2","A4","A3 Full"]))
    for key in order:
        row=summary.add_row().cells
        vals=[codes[key],pct(s.loc[key,'json_parse_mean']),pct(s.loc[key,'srs_form_mean']),pct(s.loc[key,'traceability_mean']),pct(s.loc[key,'full_compliance_mean']),f"{s.loc[key,'topic_alignment_f1_mean']:.3f}",f"{s.loc[key,'repair_count_mean']:.2f}"]
        for j,v in enumerate(vals): row[j].text=v
    style_table(summary, [1020, 1140, 1140, 1440, 1440, 1680, 1500], numeric_cols=(1,2,3,4,5,6))
    add_caption(doc, "表2　六个条件在30个冻结案例上的主要结果。")
    doc.add_picture(str(FIG / "outcome_bars.png"), width=Inches(6.45))
    add_caption(doc, "图1　结构、工程与主题对齐结果。所有条件JSON和SRS均为100%，主要差异来自精确证据绑定。")

    doc.add_page_break()
    add_heading(doc, "5.1 完整系统的配对效应", 2)
    effect = doc.add_table(rows=1, cols=7)
    for j,h in enumerate(["比较", "指标", "完整", "对照", "差值", "95% CI", "Holm p"]): effect.cell(0,j).text=h
    compare_order = [
        ("B0_single_llm","B0"),("B1_single_llm_rag","B1"),("A1_without_affect","A1"),("A2_without_validator","A2"),("A4_without_need_rag","A4")
    ]
    for cond_key, code in compare_order:
        comp = f"A3_full_agent - {cond_key}"
        for metric, name in [("traceability","Trace"),("topic_alignment_f1","Topic F1")]:
            x = c[(c.comparison==comp)&(c.metric==metric)].iloc[0]
            scale = 100 if metric=="traceability" else 1
            vals=[f"A3−{code}",name,f"{x.full_mean*scale:.3f}" if scale==1 else pct(x.full_mean),f"{x.other_mean*scale:.3f}" if scale==1 else pct(x.other_mean),f"{x.difference*scale:+.3f}" if scale==1 else f"{x.difference*100:+.1f} pp",f"[{x.ci_low*scale:.3f}, {x.ci_high*scale:.3f}]" if scale==1 else f"[{x.ci_low*100:.1f}, {x.ci_high*100:.1f}] pp",pformat(x.p_holm)]
            row=effect.add_row().cells
            for j,v in enumerate(vals): row[j].text=v
    style_table(effect, [1080, 1080, 1260, 1260, 1500, 2040, 1140], numeric_cols=(2,3,4,5,6))
    add_caption(doc, "表3　完整系统与各条件的配对差异。pp表示百分点；Holm p按指标族校正。")
    doc.add_picture(str(FIG / "traceability_effects.png"), width=Inches(6.3))
    add_caption(doc, "图2　完整系统相对各条件的Traceability配对差值与95% bootstrap区间。")

    add_heading(doc, "6. 模块级解释", 1)
    add_heading(doc, "6.1 Validator：明确的工程方向，但样本量不足", 2)
    add_body(doc, "在完全共享原始生成的严格比较中，A2（无Validator）的Traceability和Full compliance均为83.3%，A3完整系统均为100.0%，绝对提升16.7个百分点；95% bootstrap区间为[3.3, 30.0]个百分点。exact McNemar原始p=0.0625，Holm校正p=0.1875，因此不应写成统计显著。")
    add_body(doc, "Validator共执行7次 evidence quote规范化，涉及6/30个案例；没有改变Topic-alignment F1或词汇接地代理。这说明观察到的增益来自证据字段的精确规范，而不是需求语义被重新生成。")

    add_heading(doc, "6.2 Affect Interpreter：未观察到可归因增益", 2)
    add_body(doc, "A3与A1在Traceability和Full compliance上均为100.0%。Topic-alignment F1由0.625升至0.644，差值0.019，95%区间[-0.037, 0.080]，Holm p=1.000。自动指标不能证明Affect Interpreter改善了语义主题理解。若论文要保留相关主张，需要盲评“情感线索识别、过度解释和隐含需求有效性”。")

    add_heading(doc, "6.3 Need/RAG：小幅主题方向，无显著证据", 2)
    add_body(doc, "去掉Need/RAG的A4与完整A3在结构指标上完全相同。Topic-alignment F1为0.606与0.644，完整系统高0.039，但95%区间[-0.034, 0.110]，Holm p=1.000。B1（Single LLM+RAG）也没有优于B0；其Traceability反而由80.0%降至76.7%。因此，本次结果不支持将RAG单独描述为已验证有效。")

    add_heading(doc, "6.4 完整系统与单次LLM", 2)
    add_body(doc, "A3相对B0的Traceability和Full compliance均提高20.0个百分点，bootstrap区间[6.7, 33.3]；原始McNemar p=0.0313，但Holm p=0.125。Topic-alignment F1仅由0.636升至0.644，差值0.008，未显示语义优势。结果方向与论文的“工程控制与可审计性”定位一致，但证据强度应称为探索性。")

    add_heading(doc, "7. 工程成本", 1)
    cost = doc.add_table(rows=1, cols=4)
    for j,h in enumerate(["条件", "延迟(s/例)", "输入token/例", "输出token/例"]): cost.cell(0,j).text=h
    for key in order:
        row=cost.add_row().cells
        vals=[codes[key],f"{s.loc[key,'latency_s_mean']:.2f}",f"{s.loc[key,'prompt_tokens_mean']:.1f}",f"{s.loc[key,'completion_tokens_mean']:.1f}"]
        for j,v in enumerate(vals): row[j].text=v
    style_table(cost, [1800, 2280, 2640, 2640], numeric_cols=(1,2,3))
    add_caption(doc, "表4　端到端API时间和token。含Affect的条件计入独立Affect调用。")
    doc.add_picture(str(FIG / "cost_bars.png"), width=Inches(6.45))
    add_caption(doc, "图3　各条件的工程成本。")
    add_body(doc, "完整系统平均延迟12.21秒，相比B0的8.02秒增加4.19秒（约52%）；平均输入token由438.2增至1,007.6，增加约130%；输出token由346.8增至464.5，增加约34%。进入主分析的150次唯一调用共使用79,421个输入token和46,448个输出token。")

    add_heading(doc, "8. 有效性威胁与不可过度解释的部分", 1)
    limitations = [
        ("重建而非原始复现。", "原始30例、代码和提示词不可得，因此结果与论文现有数值不能直接合并。"),
        ("样本量较小。", "30例使McNemar检验对5–7个不一致案例的统计功效有限，多重校正后无比较达到0.05。"),
        ("强模型造成天花板。", "JSON Mode和较强API模型使所有条件JSON与SRS均为100%，这些指标失去区分度。"),
        ("语义有效性尚未人工确认。", "Topic F1和词汇覆盖只是自动代理，不能替代独立需求工程评价者对Evidence-based Requirement Validity的盲评。"),
        ("Validator指标存在机制耦合。", "Validator直接规范证据字段，因此Traceability提升证明工程控制有效，不等价于需求内容更正确。"),
        ("API非完全确定。", "严格共享原始输出解决了Validator归因，但其他跨模块比较仍可能混入少量服务端生成波动。"),
    ]
    lim = doc.add_table(rows=1, cols=2)
    lim.cell(0,0).text,lim.cell(0,1).text="威胁","影响"
    for a,b in limitations:
        row=lim.add_row().cells;row[0].text=a;row[1].text=b
    style_table(lim,[2520,6840])

    add_heading(doc, "9. 对论文表述的直接建议", 1)
    add_body(doc, "可以保留的结论：在本次Qwen API重建实验中，确定性Validator将精确证据绑定和完整工程合规率由83.3%提高到100.0%，显示出积极工程效应方向；完整框架相对单次LLM也呈现更高的可追溯性。")
    add_body(doc, "必须弱化的结论：不能声称Affect Interpreter或Need/RAG显著改善语义理解；不能把100% Traceability解释为需求内容100%正确；不能把未经独立评价者确认的自动代理称为Evidence-based Requirement Validity。")
    add_body(doc, "建议论文主结果措辞：『消融结果表明，完整框架的主要可观察增益来自证据规范和验证控制，而非主题语义对齐。移除Validator后，Traceability与Full compliance均下降16.7个百分点；但由于样本量为30且经多重比较校正，结果应视为探索性证据。』")
    add_body(doc, "建议后续确认实验：新增60–100个预注册测试案例；由至少2名需求工程评价者进行盲评；报告Cohen’s κ或Krippendorff’s α；将Unsupported requirement rate、Validator false rejection和semantic preservation列为主要指标。")

    add_heading(doc, "10. 结论", 1)
    p = add_body(doc, "本次重建消融完成了论文原计划但未执行的模块移除比较。最稳健的解释是：Validator改善了证据字段的确定性、规范性和审计性；Affect与Need/RAG在当前30例和自动指标下没有显示显著独立贡献；完整框架付出了更高延迟和token成本。该结果支持将论文贡献聚焦于工程控制与可追溯性，同时要求对语义和情感理解主张保持克制。")
    shade_paragraph(p, PALE)

    add_heading(doc, "附录A　可复现性资产", 1)
    artifacts = [
        ("run_ablation_api.py", "数据冻结、分类/检索、API调用、Validator、指标与统计"),
        ("protocol_config.json", "模型、条件、数据哈希、拆分和输出文件哈希"),
        ("frozen_cases.csv", "30例冻结测试清单及evidence spans"),
        ("api_responses.jsonl", "210次完整API响应、token、延迟与request ID"),
        ("predictions.jsonl", "主分析采用的逐例原始/最终输出、修复记录与指标"),
        ("condition_summary.csv", "条件级均值与标准差"),
        ("paired_comparisons.csv", "配对差值、95%区间、原始p与Holm p"),
    ]
    at = doc.add_table(rows=1,cols=2);at.cell(0,0).text="文件";at.cell(0,1).text="内容"
    for a,b in artifacts:
        row=at.add_row().cells;row[0].text=a;row[1].text=b
    style_table(at,[3000,6360])
    add_body(doc, f"数据文件SHA-256：{cfg['dataset_sha256']}。API Key未写入任何报告或结果文件；仅存放于权限受限且被.gitignore排除的本地.env文件。")

    doc.core_properties.title = "Qwen API 消融实验结果报告"
    doc.core_properties.subject = "Affect-sensitive requirement elicitation multi-agent ablation study"
    doc.core_properties.author = "Independent reconstructed experiment"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
